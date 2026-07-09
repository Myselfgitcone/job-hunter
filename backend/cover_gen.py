"""
Cover letter → PDF / DOCX.

Letter format (not resume format): candidate name header + contact line +
date + body paragraphs. Calibri 11pt in DOCX; Helvetica 10.5pt in PDF.
No tables, no text boxes — ATS/parser safe.
"""
import io
import re
from datetime import date

# ── PDF ───────────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import letter as _LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_LEFT

_BLACK = colors.HexColor("#1A1A1A")
_GRAY  = colors.HexColor("#555555")
_NAVY  = colors.HexColor("#1F3864")


def _clean_paragraphs(text: str) -> list[str]:
    """Split letter body into paragraphs on blank lines; join wrapped lines."""
    blocks: list[str] = []
    cur: list[str] = []
    for ln in (text or "").replace("\r\n", "\n").split("\n"):
        if ln.strip():
            cur.append(ln.strip())
        elif cur:
            blocks.append(" ".join(cur))
            cur = []
    if cur:
        blocks.append(" ".join(cur))
    return blocks


def _esc(t: str) -> str:
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def generate_cover_pdf(letter_text: str, candidate_name: str = "", contact_line: str = "",
                       company: str = "") -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=_LETTER,
        leftMargin=0.8 * inch, rightMargin=0.8 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
        title=f"{candidate_name} — Cover Letter" + (f" | {company}" if company else ""),
        author=candidate_name or "Candidate",
    )
    name_st = ParagraphStyle("name", fontName="Helvetica-Bold", fontSize=16,
                             leading=19, textColor=_NAVY, alignment=TA_LEFT)
    meta_st = ParagraphStyle("meta", fontName="Helvetica", fontSize=9.5,
                             leading=12, textColor=_GRAY, alignment=TA_LEFT)
    body_st = ParagraphStyle("body", fontName="Helvetica", fontSize=10.5,
                             leading=15, textColor=_BLACK, alignment=TA_LEFT,
                             spaceAfter=10)
    story = []
    if candidate_name:
        story.append(Paragraph(_esc(candidate_name), name_st))
    if contact_line:
        story.append(Spacer(1, 2))
        story.append(Paragraph(_esc(contact_line), meta_st))
    story.append(Spacer(1, 4))
    story.append(HRFlowable(width="100%", thickness=1, color=_NAVY))
    story.append(Spacer(1, 14))
    story.append(Paragraph(date.today().strftime("%B %d, %Y"), meta_st))
    story.append(Spacer(1, 12))
    for para in _clean_paragraphs(letter_text):
        story.append(Paragraph(_esc(para), body_st))
    doc.build(story)
    return buf.getvalue()


# ── DOCX ──────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_D_NAVY  = RGBColor(0x1F, 0x38, 0x64)
_D_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
_D_GRAY  = RGBColor(0x55, 0x55, 0x55)
_D_FONT  = "Calibri"


def _run(p, text, size, color, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = _D_FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def _spacing(p, before=0, after=0):
    pf = p.paragraph_format
    if before: pf.space_before = Pt(before)
    if after:  pf.space_after = Pt(after)
    pf.line_spacing = 1


def _rule(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1F3864")
    pBdr.append(bottom)
    pPr.insert(0, pBdr)


def generate_cover_docx(letter_text: str, candidate_name: str = "", contact_line: str = "",
                        company: str = "") -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.8)
    section.top_margin = section.bottom_margin = Inches(0.7)

    if candidate_name:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _spacing(p, after=2)
        _run(p, candidate_name, 16, _D_NAVY, bold=True)
    if contact_line:
        p = doc.add_paragraph(); _spacing(p, after=4)
        _rule(p)
        _run(p, contact_line, 10.5, _D_GRAY)
    p = doc.add_paragraph(); _spacing(p, before=10, after=10)
    _run(p, date.today().strftime("%B %d, %Y"), 10.5, _D_GRAY)

    for para in _clean_paragraphs(letter_text):
        p = doc.add_paragraph(); _spacing(p, after=9)
        _run(p, para, 11, _D_BLACK)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
