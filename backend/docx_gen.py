"""
Generate DOCX matching exact formatting of Jagadish_Resume_Resume.docx reference.

Specs extracted from reference:
  Name:    18pt  bold  NAVY  CENTER
  Title:   12pt        GRAY  CENTER
  Contact: 9.5pt       GRAY  CENTER  + NAVY bottom border
  Section: 11pt  bold  NAVY         + NAVY bottom border  space_before=11pt after=4pt
  JobHdr:  10.5pt bold BLACK        + right-tab for date  space_before=8pt after=2.5pt
  Bullet:  9.5pt       BLACK  List Paragraph  left=440 hanging=260 twips  after=1.6pt
  Tech:    9pt  italic  GRAY   indent left=200 twips  "Technologies Used:" bold+italic
  Skill:   9.5pt       BLACK  List Paragraph (same bullet)  bold label + plain value
  Edu:     9.5pt bold BLACK + GRAY plain
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy
import re
import io

NAVY  = RGBColor(0x1F, 0x38, 0x64)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY  = RGBColor(0x55, 0x55, 0x55)

SZ_NAME    = Pt(17)
SZ_TITLE   = Pt(11)
SZ_CONTACT = Pt(9)
SZ_SECTION = Pt(10.5)
SZ_JOB     = Pt(10)
SZ_BODY    = Pt(9)
SZ_TECH    = Pt(8.5)

FONT = "Calibri"


# ── XML helpers ───────────────────────────────────────────────────────────────

def _set_run(run, size, color, bold=None, italic=None):
    run.font.name  = FONT
    run.font.size  = size
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_spacing(para, before_pt=0, after_pt=0):
    pf = para.paragraph_format
    if before_pt:
        pf.space_before = Pt(before_pt)
    if after_pt:
        pf.space_after = Pt(after_pt)
    pf.line_spacing = 1  # single


def _add_bottom_border(para, color_hex, size=6, space=4):
    pPr = para._p.get_or_add_pPr()
    # Remove existing border if any
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.insert(0, pBdr)


def _add_right_tab(para):
    """Add a right-align tab stop at the right margin."""
    pPr = para._p.get_or_add_pPr()
    tabs_el = pPr.find(qn("w:tabs"))
    if tabs_el is None:
        tabs_el = OxmlElement("w:tabs")
        pPr.append(tabs_el)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "10656")   # 7.4" content width in twips (0.55" margins)
    tabs_el.append(tab)


def _set_indent(para, left_twips=0, hanging_twips=0):
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if left_twips:
        ind.set(qn("w:left"), str(left_twips))
    if hanging_twips:
        ind.set(qn("w:hanging"), str(hanging_twips))


def _ensure_numbering(doc):
    """
    Add a bullet numbering definition (numId=1) to the document if not present.
    Bullet char "•", left=440 twips, hanging=260 twips — exact match to reference.
    """
    # Check if numbering part exists
    try:
        num_part = doc.part.numbering_part
    except Exception:
        num_part = None

    if num_part is None:
        # Create numbering part from scratch
        from docx.opc.part import Part
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml.parser import parse_xml
        numbering_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:abstractNum w:abstractNumId="0">'
            '<w:multiLevelType w:val="hybridMultilevel"/>'
            '<w:lvl w:ilvl="0">'
            '<w:start w:val="1"/>'
            '<w:numFmt w:val="bullet"/>'
            '<w:lvlText w:val="&#x2022;"/>'
            '<w:lvlJc w:val="left"/>'
            '<w:pPr><w:ind w:left="440" w:hanging="260"/></w:pPr>'
            '<w:rPr>'
            '<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
            '<w:color w:val="1A1A1A"/>'
            '<w:sz w:val="19"/>'
            '</w:rPr>'
            '</w:lvl>'
            '</w:abstractNum>'
            '<w:num w:numId="1">'
            '<w:abstractNumId w:val="0"/>'
            '</w:num>'
            '</w:numbering>'
        )
        part = Part(
            "/word/numbering.xml",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            parse_xml(numbering_xml.encode()),
            doc.part.package,
        )
        doc.part.relate_to(part, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering")
        return

    # Numbering part exists — ensure our abstractNum/num exists (numId=1)
    numbering_el = num_part._element
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    existing = numbering_el.findall(f"{{{ns}}}abstractNum")
    if existing:
        return  # Already has numbering defs

    abstract_xml = (
        f'<w:abstractNum xmlns:w="{ns}" w:abstractNumId="0">'
        '<w:multiLevelType w:val="hybridMultilevel"/>'
        '<w:lvl w:ilvl="0">'
        '<w:start w:val="1"/>'
        '<w:numFmt w:val="bullet"/>'
        '<w:lvlText w:val="&#x2022;"/>'
        '<w:lvlJc w:val="left"/>'
        '<w:pPr><w:ind w:left="440" w:hanging="260"/></w:pPr>'
        '<w:rPr>'
        '<w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
        '<w:color w:val="1A1A1A"/>'
        '<w:sz w:val="19"/>'
        '</w:rPr>'
        '</w:lvl>'
        '</w:abstractNum>'
    )
    num_xml = (
        f'<w:num xmlns:w="{ns}" w:numId="1">'
        '<w:abstractNumId w:val="0"/>'
        '</w:num>'
    )
    numbering_el.insert(0, etree.fromstring(abstract_xml))
    numbering_el.append(etree.fromstring(num_xml))


def _apply_bullet_numbering(para):
    """Wire paragraph to numId=1, level 0."""
    pPr = para._p.get_or_add_pPr()
    # Remove existing numPr
    for old in pPr.findall(qn("w:numPr")):
        pPr.remove(old)
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.insert(0, numPr)


# ── Paragraph builders ────────────────────────────────────────────────────────

def _add_name(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before_pt=0, after_pt=3)
    r = p.add_run(text)
    _set_run(r, SZ_NAME, NAVY, bold=True)
    return p


def _add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before_pt=0, after_pt=3)
    r = p.add_run(text)
    _set_run(r, SZ_TITLE, GRAY)
    return p


def _add_contact(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, before_pt=0, after_pt=4)
    _add_bottom_border(p, "1F3864", size=8, space=4)
    r = p.add_run(text)
    _set_run(r, SZ_CONTACT, GRAY)
    return p


def _add_section(doc, text):
    p = doc.add_paragraph()
    _set_spacing(p, before_pt=7, after_pt=2)
    _add_bottom_border(p, "1F3864", size=6, space=3)
    r = p.add_run(text)
    _set_run(r, SZ_SECTION, NAVY, bold=True)
    return p


def _add_job_header(doc, line):
    """
    Parse: "Title @ Company | Location    Date"
    or:    "Title @ Company | Location    Date"  (— separated)
    """
    p = doc.add_paragraph()
    _set_spacing(p, before_pt=5, after_pt=1.5)
    _add_right_tab(p)

    # Split on first " @ "
    if " @ " in line:
        before_at, after_at = line.split(" @ ", 1)
        title_text = before_at.strip()
        rest = after_at  # "Company | Location    Date"
    else:
        title_text = line
        rest = ""

    # Parse company + location/date
    company_text = rest
    loc_date_text = ""
    if " | " in rest:
        company_text, loc_date = rest.split(" | ", 1)
        # Date is the last token after multiple spaces or em-dash
        # Try to detect date pattern like "Sep 2023 – Present" or "Jan 2021 – Jul 2022"
        date_match = re.search(
            r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}.*$)",
            loc_date
        )
        if date_match:
            location = loc_date[:date_match.start()].strip()
            date = date_match.group(1).strip()
            loc_date_text = f"  |  {location}\t{date}"
        else:
            loc_date_text = f"  |  {loc_date}"

    # Build runs
    r1 = p.add_run(title_text)
    _set_run(r1, SZ_JOB, BLACK, bold=True)

    if company_text:
        r2 = p.add_run(f" @ {company_text.strip()}")
        _set_run(r2, SZ_JOB, BLACK, bold=False)

    if loc_date_text:
        r3 = p.add_run(loc_date_text)
        _set_run(r3, SZ_JOB, GRAY, bold=False)

    return p


def _add_bullet(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    _apply_bullet_numbering(p)
    _set_spacing(p, after_pt=1)
    # Clear any paragraph-level indent (let numbering define it)
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:ind")):
        pPr.remove(old)
    r = p.add_run(text)
    _set_run(r, SZ_BODY, BLACK)
    return p


def _add_tech(doc, rest_text):
    p = doc.add_paragraph()
    _set_spacing(p, before_pt=1, after_pt=2)
    _set_indent(p, left_twips=200)
    r1 = p.add_run("Technologies Used: ")
    _set_run(r1, SZ_TECH, GRAY, bold=True, italic=True)
    r2 = p.add_run(rest_text)
    _set_run(r2, SZ_TECH, GRAY, italic=True)
    return p


def _add_skill(doc, label, value):
    p = doc.add_paragraph(style="List Paragraph")
    _apply_bullet_numbering(p)
    _set_spacing(p, after_pt=1)
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:ind")):
        pPr.remove(old)
    r1 = p.add_run(f"{label}: ")
    _set_run(r1, SZ_BODY, BLACK, bold=True)
    r2 = p.add_run(value)
    _set_run(r2, SZ_BODY, BLACK, bold=False)
    return p


def _add_education(doc, line):
    p = doc.add_paragraph()
    _set_spacing(p, before_pt=2, after_pt=2)
    # "Degree — University" or "Degree @ University"
    sep = " — " if " — " in line else (" @ " if " @ " in line else None)
    if sep:
        degree, uni = line.split(sep, 1)
        r1 = p.add_run(degree.strip())
        _set_run(r1, SZ_BODY, BLACK, bold=True)
        r2 = p.add_run(f"   —   {uni.strip()}")
        _set_run(r2, SZ_BODY, GRAY, bold=False)
    else:
        r = p.add_run(line)
        _set_run(r, SZ_BODY, BLACK, bold=True)
    return p


# ── Main generator ────────────────────────────────────────────────────────────

def generate_docx(resume_text: str, job_title: str = "", company: str = "") -> bytes:
    doc = Document()

    # Page: US Letter, 0.7" margins
    section = doc.sections[0]
    section.page_width    = Twips(12240)
    section.page_height   = Twips(15840)
    section.left_margin      = Inches(0.55)
    section.right_margin     = Inches(0.55)
    section.top_margin       = Inches(0.55)
    section.bottom_margin    = Inches(0.55)
    section.header_distance  = Inches(0.4)
    section.footer_distance  = Inches(0.4)

    # Ensure bullet numbering available
    _ensure_numbering(doc)

    lines = resume_text.strip().split("\n")

    # ── Header block ─────────────────────────────────────────────────────────
    name_line    = lines[0].strip() if lines else ""
    contact_line = lines[1].strip() if len(lines) > 1 else ""

    name_parts = name_line.split(" — ", 1)
    name_text  = name_parts[0].strip()
    title_text = name_parts[1].strip() if len(name_parts) > 1 else ""

    _add_name(doc, name_text)
    if title_text:
        _add_title(doc, title_text)
    _add_contact(doc, contact_line)

    # ── Body ─────────────────────────────────────────────────────────────────
    in_skills   = False
    in_education = False

    for line in lines[2:]:
        line = line.rstrip()

        if not line:
            continue  # skip blank lines — spacing handled by para spacing

        # Section headers: ALL CAPS ending with ":"
        if (line == line.upper() and len(line) > 3
                and line.endswith(":") and not line.startswith("•")):
            section_name = line.rstrip(":")
            in_skills    = "SKILL" in section_name or "TECHNICAL" in section_name
            in_education = "EDUC" in section_name
            _add_section(doc, section_name)
            continue

        # Education section body
        if in_education:
            _add_education(doc, line)
            continue

        # Technologies Used
        if line.startswith("Technologies Used:"):
            rest = line[len("Technologies Used:"):].strip()
            _add_tech(doc, rest)
            continue

        # Bullets
        if line.startswith("•"):
            text = line[1:].strip()
            if in_skills and ":" in text:
                # Skill line: "Label: value"
                label, _, value = text.partition(":")
                _add_skill(doc, label.strip(), value.strip())
            else:
                _add_bullet(doc, text)
            continue

        # Job title lines: "Title @ Company | Location  Date"
        if re.match(r"^.+? @ .+", line):
            _add_job_header(doc, line)
            continue

        # Default body
        p = doc.add_paragraph()
        _set_spacing(p, after_pt=2)
        r = p.add_run(line)
        _set_run(r, SZ_BODY, BLACK)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
