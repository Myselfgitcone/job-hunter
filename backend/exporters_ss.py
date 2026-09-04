"""
StackShift-style PDF/DOCX export.

Replicates the reference resume's typography exactly:

  Page       : US Letter, margins  L/R 0.40"  T/B 0.50"
  Font       : Arial (embedded when available; Helvetica core fallback)
  Color      : 100% black
  Leading    : 13 pt
  Name+Title : 13 pt bold, one line  "Name - Title"
  Contact    : 11 pt regular, phone first then email
  Section hdr: 12 pt bold, UPPERCASE + ':', with a 0.8 pt full-width rule under
  Body/bullet: 11 pt regular, bullet = •
  Fit        : never truncates — uniformly shrinks type until it fits 2 pages

Input is job-hunter's plain-text resume format (UPPERCASE section headers with
a colon, "•" bullets, "Title @ Company | Location Dates" job headers). Markdown
input is also accepted so the classifier keeps working if a draft slips.

Drop-in replacement for pdf_gen.generate_pdf / docx_gen.generate_docx —
same signatures.
"""
import io
import os
import re
from html import escape

BLACK = "#000000"

# ---- sizes / spacing (points) ---------------------------------------------
BODY = 11.0
NAME = 13.0          # name + title = body + 2, bold black
LEADING = 13.0
MARGIN_LR = 0.40 * 72
MARGIN_TB = 0.50 * 72
RULE_WIDTH = 0.8

# ---------------------------------------------------------------------------
# Font registration (Arial when we can find it; Helvetica core otherwise)
# ---------------------------------------------------------------------------
_FONT = "Helvetica"
_FONT_B = "Helvetica-Bold"
_DOCX_FONT = "Arial"     # Word resolves this on the reader's machine regardless
_fonts_ready = False


def _ensure_fonts():
    global _FONT, _FONT_B, _fonts_ready
    if _fonts_ready:
        return
    _fonts_ready = True
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        win = os.environ.get("WINDIR", "C:/Windows")
        bundled = os.path.join(os.path.dirname(__file__), "fonts")
        # Bundled TTFs first so Linux/Railway matches local output; then the
        # Windows copy of Arial; then common Linux Liberation/DejaVu paths.
        candidates = [
            (os.path.join(bundled, "Arial.ttf"), os.path.join(bundled, "Arial-Bold.ttf")),
            (os.path.join(bundled, "LiberationSans-Regular.ttf"),
             os.path.join(bundled, "LiberationSans-Bold.ttf")),
            (os.path.join(win, "Fonts", "arial.ttf"), os.path.join(win, "Fonts", "arialbd.ttf")),
            ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
             "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"),
            ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
             "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ]
        # Nix installs fonts under an unpredictable /nix/store hash, so glob for
        # them rather than hardcoding a path.
        try:
            import glob as _glob
            for pat_r, pat_b in (
                ("/nix/store/*/share/fonts/**/LiberationSans-Regular.ttf",
                 "/nix/store/*/share/fonts/**/LiberationSans-Bold.ttf"),
                ("/usr/share/fonts/**/LiberationSans-Regular.ttf",
                 "/usr/share/fonts/**/LiberationSans-Bold.ttf"),
                ("/usr/share/fonts/**/DejaVuSans.ttf",
                 "/usr/share/fonts/**/DejaVuSans-Bold.ttf"),
            ):
                hits_r = _glob.glob(pat_r, recursive=True)
                hits_b = _glob.glob(pat_b, recursive=True)
                if hits_r and hits_b:
                    candidates.append((hits_r[0], hits_b[0]))
        except Exception:  # noqa: BLE001 — globbing is opportunistic
            pass

        for reg, bold in candidates:
            if os.path.exists(reg) and os.path.exists(bold):
                pdfmetrics.registerFont(TTFont("Arial", reg))
                pdfmetrics.registerFont(TTFont("Arial-Bold", bold))
                # Map the family so inline <b> markup resolves to the bold face.
                pdfmetrics.registerFontFamily(
                    "Arial", normal="Arial", bold="Arial-Bold",
                    italic="Arial", boldItalic="Arial-Bold",
                )
                _FONT, _FONT_B = "Arial", "Arial-Bold"
                print(f"[EXPORT] Embedded font: {reg}")
                break
        else:
            print("[EXPORT] No TTF found — using Helvetica core (metrically close to Arial)")
    except Exception as exc:  # noqa: BLE001 — keep the Helvetica fallback
        print(f"[EXPORT] Font registration failed ({exc}) — using Helvetica")


# ---------------------------------------------------------------------------
# Text -> typed elements
# ---------------------------------------------------------------------------
_BOLD = re.compile(r"\*\*(.+?)\*\*")
_BOLD_LINE = re.compile(r"^\*\*(.+?)\*\*[\s,–—-]*$")
_TECHLABEL = re.compile(r"^[-*•]?\s*\*{0,2}(technolog|tech stack|tools used|environment)", re.I)
_CONTACT = re.compile(r"[@|]|linkedin|github|https?://|\(\d{3}\)|\d{3}[.\-]\d{3}[.\-]\d{4}", re.I)
_BULLET_START = re.compile(r"^[•\-*]\s+")

# "Sep 2024 – Present", "Jan 2021 - Jul 2022", "2019 — 2021"
_MONTHS = "Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec"
_DATE_RANGE = re.compile(
    rf"((?:{_MONTHS})?\.?\s*\d{{4}}\s*[–—\-]+\s*(?:Present|Current|(?:{_MONTHS})?\.?\s*\d{{4}}))",
    re.I,
)
_DATE_ANY = re.compile(r"((?:19|20)\d{2}|present)", re.I)
# Section header: UPPERCASE, optional trailing colon
_SECTION_LINE = re.compile(r"^[A-Z][A-Z &/,\-']{3,}:?$")


def _clean_contact(text: str) -> str:
    """Phone first, email second, then a City, ST if present; drop links and
    street addresses (recruiter platforms filter by location, so the city
    stays — nothing more personal than that goes on the page)."""
    parts = [p.strip() for p in re.split(r"[|•]", text) if p.strip()]
    phone = next((p for p in parts if re.search(r"\d{3}[)\s.\-]*\d{3}[\s.\-]*\d{4}", p)), None)
    email = next((p for p in parts if "@" in p), None)
    # "Maryland Heights, MO" / "Austin, Texas": has a comma, no digits, no URL.
    city = next((p for p in parts
                 if "," in p and not re.search(r"\d|@|http|linkedin|github|www\.", p, re.I)
                 and len(p) <= 40), None)
    keep = [x for x in (phone, email, city) if x]
    return " | ".join(keep) if keep else text


def _split_jobline(text: str):
    """Split a job header into (left, right-date).

    Handles job-hunter's shape, where the date trails the location inside the
    same pipe field ("... | Minneapolis, MN Sep 2024 – Present"), as well as
    the date living in its own field.
    """
    m = _DATE_RANGE.search(text)
    if m:
        left = text[: m.start()].rstrip().rstrip("|,").rstrip()
        return left, m.group(1).strip()

    # Fall back to per-field scanning; drop placeholder fields while we're here.
    fields = [
        f.strip() for f in text.split("|")
        if not re.fullmatch(
            r"(location\s*)?(not listed|not specified|n/?a|none|unknown|tbd)", f.strip(), re.I)
    ]
    for i, f in enumerate(fields):
        if _DATE_ANY.search(f):
            left = " | ".join(fields[:i] + fields[i + 1:])
            return left.strip(" |"), fields[i]
    return text, ""


def _demojibake(text: str) -> str:
    """Repair UTF-8-decoded-as-Latin-1 artifacts (e.g. "â€¢" → "•", "â€™" → "'")
    that creep in from resume upload parsing, so PDF/DOCX render clean glyphs."""
    if "â€" not in text and "Ã" not in text:
        return text
    for bad, good in (
        ("â€¢", "•"), ("â€“", "–"), ("â€”", "—"),
        ("â€™", "’"), ("â€˜", "‘"),
        ("â€œ", "“"), ("â€\x9d", "”"), ("â€", "”"),
        ("Â ", " "), ("Â", ""),
    ):
        text = text.replace(bad, good)
    return text


def _classify(text: str):
    """Type each line. Accepts plain text (primary) and markdown (fallback)."""
    text = _demojibake(text)
    out = []
    name_done = False
    section_seen = False
    in_skills = False
    prev_kind = None

    for raw in text.replace("\r\n", "\n").split("\n"):
        s = raw.strip()
        if not s:
            prev_kind = None
            continue
        if re.fullmatch(r"[-*_]{3,}", s):
            continue  # horizontal rule
        if s.startswith(">") or re.search(r"see above|consolidated under", s, re.I):
            continue  # consolidation stub

        # ── markdown fallbacks ────────────────────────────────────────────
        if s.startswith("## ") or s.startswith("### "):
            label = s.lstrip("#").strip()
            in_skills = "skill" in label.lower()
            out.append(("section", label))
            section_seen = True
        elif s.startswith("# "):
            out.append(("name", s[2:].strip()))
            name_done = True

        # ── plain-text format (what the tailor engine emits) ──────────────
        elif _SECTION_LINE.match(s) and len(s) < 60:
            label = s.rstrip(":")
            in_skills = "skill" in label.lower()
            out.append(("section", label))
            section_seen = True
        elif _TECHLABEL.match(s):
            out.append(("tech", _BULLET_START.sub("", s).replace("*", "").strip()))
        elif _BULLET_START.match(s):
            body = _BULLET_START.sub("", s).replace("*", "").strip()
            out.append(("skillbullet" if in_skills else "bullet", body))
        elif _BOLD_LINE.match(s):
            out.append(("jobtitle" if section_seen else "headline",
                        _BOLD_LINE.match(s).group(1).strip()))
        elif not name_done and not section_seen:
            out.append(("name", s.replace("*", "").strip()))
            name_done = True
        elif prev_kind == "name" and _CONTACT.search(s):
            out.append(("contact", s))
        elif section_seen and (" @ " in s or _DATE_RANGE.search(s)):
            out.append(("jobtitle", s.replace("*", "").strip()))
        else:
            out.append(("body", s))
        prev_kind = out[-1][0] if out else None

    return out


def _split_name_headline(name: str):
    """'Jagadish Butukuri — Senior Data Engineer' -> (name, headline)."""
    for sep in ("—", " - ", " – ", " | "):
        if sep in name:
            left, _, right = name.partition(sep)
            if left.strip() and right.strip():
                return left.strip(), right.strip()
    return name.strip(), None


def _label_markup(text: str) -> str:
    """Bold the label up to the first colon, regular the rest."""
    text = text.replace("*", "").strip()
    if ":" in text:
        label, rest = text.split(":", 1)
        return f"<b>{escape(label)}:</b>{escape(rest)}"
    return escape(text)


def _inline(text: str) -> str:
    parts, pos = [], 0
    for m in _BOLD.finditer(text):
        parts.append(escape(text[pos:m.start()]))
        parts.append("<b>" + escape(m.group(1)) + "</b>")
        pos = m.end()
    parts.append(escape(text[pos:]))
    return "".join(parts)


_CONTENT_KINDS = {"bullet", "skillbullet", "tech", "body"}


def _prune_empty(elements):
    """Drop bulletless job stubs, then section headers left with nothing."""
    # Pass 1: in Experience/Projects a job title needs at least one bullet.
    # Education/Certifications titles are legitimately bulletless.
    keep = [True] * len(elements)
    cur_sec = ""
    for i, (kind, text) in enumerate(elements):
        if kind == "section":
            cur_sec = text.lower()
            continue
        if kind != "jobtitle" or not ("experience" in cur_sec or "project" in cur_sec):
            continue
        has = False
        for j in range(i + 1, len(elements)):
            k2 = elements[j][0]
            if k2 in ("jobtitle", "section"):
                break
            if k2 in _CONTENT_KINDS:
                has = True
                break
        if not has:
            keep[i] = False
    elements = [e for e, k in zip(elements, keep) if k]

    # Pass 2: a section needs ANY element before the next section.
    keep = [True] * len(elements)
    for i, (kind, _) in enumerate(elements):
        if kind != "section":
            continue
        nxt = elements[i + 1][0] if i + 1 < len(elements) else "section"
        if nxt == "section":
            keep[i] = False
    return [e for e, k in zip(elements, keep) if k]


def _split_header(elements):
    name = contact = headline = None
    rest = []
    for kind, text in elements:
        if kind == "name" and name is None:
            name, split_headline = _split_name_headline(text)
            if split_headline and headline is None:
                headline = split_headline
        elif kind == "contact" and contact is None:
            contact = text
        elif kind == "headline" and headline is None:
            headline = text
        else:
            rest.append((kind, text))
    return name, contact, headline, rest


# ---------------------------------------------------------------------------
# PDF — auto-fit: uniformly shrink type (never truncate) to cap at MAX_PAGES
# ---------------------------------------------------------------------------
TARGET_PAGES = 2          # ideal length
MAX_PAGES = 3             # acceptable when full skill coverage needs the room
_MIN_TARGET_SCALE = 0.92  # shrink harder than this to force 2 pages = cramped;
                          # a readable 3rd page beats squeezed type
_SCALES = [1.0, 0.96, 0.92, 0.88, 0.84, 0.80]

# ReportLab's bulletText glyph for U+2022 carries no usable ToUnicode mapping
# (even with an embedded TTF), so PDF text extractors read every bullet as
# "(cid:127)" garbage. The middle dot maps cleanly; rendered bold and slightly
# larger it looks like a normal bullet.
_BULLET_GLYPH = "·"


def _render_pdf(elements, scale, title="Tailored Resume"):
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Table, TableStyle

    black = HexColor(BLACK)
    body = BODY * scale
    name_sz = body + 2
    lead = LEADING * scale
    mlr = MARGIN_LR * (0.9 + 0.1 * scale)
    mtb = MARGIN_TB * (0.9 + 0.1 * scale)
    usable_w = letter[0] - 2 * mlr
    sp = scale  # spacing multiplier

    def ps(nm, size, bold=False, **kw):
        params = {"fontName": _FONT_B if bold else _FONT, "fontSize": size,
                  "leading": lead, "textColor": black, "alignment": TA_LEFT}
        params.update(kw)
        return ParagraphStyle(nm, **params)

    st_name = ps("name", name_sz, bold=True, leading=name_sz + 2, spaceAfter=1)
    st_contact = ps("contact", body, spaceAfter=4 * sp)
    st_section = ps("section", body + 1, bold=True, spaceBefore=8 * sp, spaceAfter=2 * sp)
    st_job = ps("job", body, bold=True, spaceBefore=4 * sp, spaceAfter=1)
    st_job_r = ps("jobr", body, bold=True, alignment=TA_RIGHT, spaceBefore=4 * sp, spaceAfter=1)
    st_body = ps("body", body, spaceAfter=2 * sp)
    st_tech = ps("tech", body, spaceAfter=3 * sp)
    st_bullet = ps("bullet", body, leftIndent=13, bulletIndent=2, spaceAfter=2 * sp,
                   bulletFontName=_FONT_B, bulletFontSize=body + 3)

    name, contact, headline, rest = _split_header(elements)
    story = []
    if name or headline:
        head = "<b>" + _inline(name or "") + "</b>"
        if headline:
            head += " &mdash; <b>" + _inline(headline) + "</b>"
        story.append(Paragraph(head, st_name))
    if contact:
        story.append(Paragraph(escape(_clean_contact(contact)), st_contact))

    for kind, text in rest:
        if kind == "section":
            story.append(Paragraph(escape(text).upper().rstrip(":") + ":", st_section))
            story.append(HRFlowable(width="100%", thickness=RULE_WIDTH,
                                    color=black, spaceBefore=1, spaceAfter=4 * sp))
        elif kind == "jobtitle":
            left, date = _split_jobline(text)
            if date:
                lp = Paragraph("<b>" + _inline(left) + "</b>", st_job)
                rp = Paragraph("<b>" + escape(date) + "</b>", st_job_r)
                t = Table([[lp, rp]], colWidths=[usable_w * 0.72, usable_w * 0.28])
                t.hAlign = "LEFT"  # align with section headers, not the frame edge
                t.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 3 * sp),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                ]))
                story.append(t)
            else:
                story.append(Paragraph("<b>" + _inline(text) + "</b>", st_job))
        elif kind == "tech":
            story.append(Paragraph(_label_markup(text), st_tech))
        elif kind == "skillbullet":
            story.append(Paragraph(_label_markup(text), st_bullet, bulletText=_BULLET_GLYPH))
        elif kind == "bullet":
            story.append(Paragraph(_inline(text), st_bullet, bulletText=_BULLET_GLYPH))
        else:
            story.append(Paragraph(_inline(text), st_body))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, leftMargin=mlr, rightMargin=mlr,
                            topMargin=mtb, bottomMargin=mtb, title=title)
    doc.build(story)
    return buf.getvalue(), doc.page


def _choose_scale(elements, title="Tailored Resume"):
    """2 pages at readable scale when possible; otherwise 3 pages at the most
    readable scale that achieves it; tightest render as a last resort."""
    last = None
    three = None
    for scale in _SCALES:
        data, pages = _render_pdf(elements, scale, title)
        if pages <= TARGET_PAGES and scale >= _MIN_TARGET_SCALE:
            if scale < 0.95:
                print(f"[EXPORT] WARNING dense page: type scaled to {scale} to fit "
                      f"{pages} page(s) — the tailor length guard should have trimmed")
            elif scale < 1.0:
                print(f"[EXPORT] PDF fit at scale {scale} ({pages} page(s))")
            return scale, data
        if pages <= MAX_PAGES and three is None:
            three = (scale, data)
        last = (scale, data)
    if three is not None:
        print(f"[EXPORT] PDF at {MAX_PAGES} pages, scale {three[0]} — content needs the room")
        return three
    print(f"[EXPORT] PDF still over {MAX_PAGES} pages at tightest scale — shipping tightest")
    return last


# ---------------------------------------------------------------------------
# Public API — drop-in for pdf_gen.generate_pdf / docx_gen.generate_docx
# ---------------------------------------------------------------------------
def generate_pdf(resume_text: str, job_title: str = "", company: str = "") -> bytes:
    _ensure_fonts()
    elements = _prune_empty(_classify(resume_text or ""))
    name = _split_header(elements)[0] or "Resume"
    title = f"{name} — Resume" + (f" | {company}" if company else "")
    return _choose_scale(elements, title)[1]


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def _para_bottom_border(paragraph):
    """Thin bottom border — the section rule."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")      # 0.75 pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _split_bold(text: str):
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            yield text[pos:m.start()], False
        yield m.group(1), True
        pos = m.end()
    if pos < len(text):
        yield text[pos:], False


def generate_docx(resume_text: str, job_title: str = "", company: str = "") -> bytes:
    _ensure_fonts()
    import docx
    from docx.shared import Inches, Pt, RGBColor

    elements = _prune_empty(_classify(resume_text or ""))
    scale = _choose_scale(elements)[0]  # same auto-fit scale as the PDF
    body = BODY * scale
    name_sz = body + 2
    lead = LEADING * scale

    black = RGBColor(0, 0, 0)
    doc = docx.Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.40)
    sec.top_margin = sec.bottom_margin = Inches(0.50)
    normal = doc.styles["Normal"]
    normal.font.name = _DOCX_FONT
    normal.font.size = Pt(body)
    normal.paragraph_format.space_after = Pt(2 * scale)
    normal.paragraph_format.line_spacing = Pt(lead)

    def runs(p, text, size, bold=False):
        for chunk, b in _split_bold(text):
            r = p.add_run(chunk)
            r.font.name = _DOCX_FONT
            r.font.size = Pt(size)
            r.bold = bold or b
            r.font.color.rgb = black

    name, contact, headline, rest = _split_header(elements)

    if name or headline:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        runs(p, name or "", name_sz, bold=True)
        if headline:
            runs(p, " — " + headline, name_sz, bold=True)
    if contact:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4 * scale)
        runs(p, _clean_contact(contact), body)

    for kind, text in rest:
        if kind == "section":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8 * scale)
            p.paragraph_format.space_after = Pt(3 * scale)
            runs(p, text.upper().rstrip(":") + ":", body + 1, bold=True)
            _para_bottom_border(p)
        elif kind == "jobtitle":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4 * scale)
            left, date = _split_jobline(text)
            if date:
                from docx.enum.text import WD_TAB_ALIGNMENT
                from docx.shared import Inches as _In
                right_pos = _In(8.5 - 0.40 - 0.40)  # page width minus L/R margins
                p.paragraph_format.tab_stops.add_tab_stop(right_pos, WD_TAB_ALIGNMENT.RIGHT)
                runs(p, left, body, bold=True)
                runs(p, "\t" + date, body, bold=True)
            else:
                runs(p, text, body, bold=True)
        elif kind in ("tech", "skillbullet"):
            p = (doc.add_paragraph(style="List Bullet") if kind == "skillbullet"
                 else doc.add_paragraph())
            p.paragraph_format.space_after = Pt((2 if kind == "skillbullet" else 3) * scale)
            t = text.replace("*", "")
            if ":" in t:
                label, restt = t.split(":", 1)
                runs(p, label + ":", body, bold=True)
                runs(p, restt, body)
            else:
                runs(p, t, body)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2 * scale)
            runs(p, text, body)
        else:
            runs(doc.add_paragraph(), text, body)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
