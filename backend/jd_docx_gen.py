"""
Generate a formatted .docx from raw JD text.
Mirrors the JdRenderer logic in the frontend.
"""
import re
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

NAVY  = RGBColor(0x1F, 0x38, 0x64)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY  = RGBColor(0x55, 0x55, 0x55)
FONT  = "Calibri"

SECTION_HEADERS = [
    "Overview", "Responsibilities", "Contributions", "Qualifications",
    "Requirements", "Required Skills", "Skills", "About", "Benefits",
    "Work Environment", "Physical Demands", "Work Authorization",
    "Equal Employment", "What You", "Who You", "What We", "Why Join",
    "Nice to Have", "Preferred", "Education", "Experience", "Summary",
    "Description", "Duties", "Job Summary", "Job Description",
    "Key Responsibilities", "Data Pipeline", "Business Intelligence",
    "Mentorship", "Collaboration", "Conditions of Employment",
    "Knowledge", "Time Type", "Compensation", "Requests",
]


def _normalize(raw: str) -> str:
    t = raw
    # bullet markers after non-newline text
    t = re.sub(r"([^\n])\s*(\* )", r"\1\n\2", t)
    t = re.sub(r"([^\n])\s*(- (?=[A-Z]))", r"\1\n\2", t)
    # section headers
    for h in SECTION_HEADERS:
        t = re.sub(r"([^\n])(" + re.escape(h) + r")", r"\1\n\n\2", t)
    # collapse 3+ newlines
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t


def _set_run(run, size, color, bold=False, italic=False):
    run.font.name  = FONT
    run.font.size  = size
    run.font.color.rgb = color
    run.bold   = bold
    run.italic = italic


def _ensure_numbering(doc):
    try:
        num_part = doc.part.numbering_part
    except Exception:
        num_part = None

    if num_part is None:
        from docx.opc.part import Part
        from docx.oxml.parser import parse_xml
        xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:abstractNum w:abstractNumId="0">'
            '<w:multiLevelType w:val="hybridMultilevel"/>'
            '<w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="bullet"/>'
            '<w:lvlText w:val="&#x2022;"/><w:lvlJc w:val="left"/>'
            '<w:pPr><w:ind w:left="440" w:hanging="260"/></w:pPr>'
            '<w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
            '<w:color w:val="1A1A1A"/><w:sz w:val="18"/></w:rPr>'
            '</w:lvl></w:abstractNum>'
            '<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
            '</w:numbering>'
        )
        part = Part(
            "/word/numbering.xml",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
            parse_xml(xml.encode()),
            doc.part.package,
        )
        doc.part.relate_to(part, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering")
        return

    ns  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    nel = num_part._element
    if not nel.findall(f"{{{ns}}}abstractNum"):
        a = (
            f'<w:abstractNum xmlns:w="{ns}" w:abstractNumId="0">'
            '<w:multiLevelType w:val="hybridMultilevel"/>'
            '<w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="bullet"/>'
            '<w:lvlText w:val="&#x2022;"/><w:lvlJc w:val="left"/>'
            '<w:pPr><w:ind w:left="440" w:hanging="260"/></w:pPr>'
            '<w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
            '<w:color w:val="1A1A1A"/><w:sz w:val="18"/></w:rPr>'
            '</w:lvl></w:abstractNum>'
        )
        n = f'<w:num xmlns:w="{ns}" w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
        nel.insert(0, etree.fromstring(a))
        nel.append(etree.fromstring(n))


def _apply_bullet(para):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:numPr")):
        pPr.remove(old)
    numPr = OxmlElement("w:numPr")
    ilvl  = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "1")
    numPr.append(ilvl); numPr.append(numId)
    pPr.insert(0, numPr)


def _add_section_header(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after  = Pt(3)
    pf.line_spacing = 1
    # bottom border
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "1F3864")
    pBdr.append(bottom); pPr.insert(0, pBdr)
    r = p.add_run(text)
    _set_run(r, Pt(11), NAVY, bold=True)


def _add_bullet_para(doc, text):
    p = doc.add_paragraph(style="List Paragraph")
    _apply_bullet(p)
    pf = p.paragraph_format
    pf.space_after  = Pt(1)
    pf.line_spacing = 1
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:ind")):
        pPr.remove(old)
    r = p.add_run(text)
    _set_run(r, Pt(9.5), BLACK)


def _add_body(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after  = Pt(3)
    pf.line_spacing = 1
    r = p.add_run(text)
    _set_run(r, Pt(9.5), BLACK)


def _is_section_header(line: str) -> bool:
    if len(line) > 80 or re.match(r"^[*\-•]", line):
        return False
    # ALL CAPS or Title Case short line
    if re.match(r"^[A-Z][A-Za-z0-9 &\/\-:,()]+$", line) and len(line) < 70:
        return True
    # Ends with colon → likely header
    if line.endswith(":") and len(line) < 60:
        return True
    return False


def generate_jd_docx(jd_text: str, company: str = "") -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.page_width   = Twips(12240)
    section.page_height  = Twips(15840)
    section.left_margin  = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin   = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    _ensure_numbering(doc)

    # Title
    if company:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"{company} — Job Description")
        _set_run(r, Pt(14), NAVY, bold=True)

    lines = _normalize(jd_text).split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Markdown heading
        if re.match(r"^#{1,3}\s", line):
            _add_section_header(doc, re.sub(r"^#+\s*", "", line))
            i += 1
            continue

        # Section header
        if _is_section_header(line):
            _add_section_header(doc, line.rstrip(":"))
            i += 1
            continue

        # Bullet group
        if re.match(r"^[*\-•]\s", line):
            while i < len(lines) and re.match(r"^[*\-•]\s", lines[i].strip()):
                text = lines[i].strip().lstrip("*-• ").strip()
                _add_bullet_para(doc, text)
                i += 1
            continue

        # Body paragraph
        _add_body(doc, line)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
